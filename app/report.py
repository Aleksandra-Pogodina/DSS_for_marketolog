"""Генерация отчёта в DOCX.

DOCX-отчёт остаётся основным стабильным форматом.
Модуль получает уже построенные графики из app.charts и вставляет их в документ.
"""

from __future__ import annotations

from datetime import datetime

import pandas as pd

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt

from app.charts import Chart
from app.processing import AnalysisResult, format_metrics_for_display


def _set_cell_bg(cell, color_hex: str) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color_hex)
    tcPr.append(shd)


def _add_dataframe_table(doc: Document, df: pd.DataFrame) -> None:
    if df is None or df.empty:
        return

    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = "Light Grid Accent 1"

    hdr = table.rows[0].cells
    for i, col in enumerate(df.columns):
        hdr[i].text = str(col)
        for paragraph in hdr[i].paragraphs:
            for run in paragraph.runs:
                run.bold = True
        _set_cell_bg(hdr[i], "3A6B8A")

    for _, row in df.iterrows():
        cells = table.add_row().cells
        for i, col in enumerate(df.columns):
            cells[i].text = "" if pd.isna(row[col]) else str(row[col])


def _format_extra_summary(result: AnalysisResult) -> pd.DataFrame:
    es = result.extra_summary.copy()
    rename = {}

    if result.extra_category_col and result.extra_category_col in es.columns:
        rename[result.extra_category_col] = str(result.extra_category_col)

    for raw, label in result.metric_labels.items():
        if raw in es.columns:
            rename[raw] = label

    if "cpc_avg" in es.columns and "cpc_avg" not in rename:
        rename["cpc_avg"] = "Средний CPC (из данных)"

    es = es.rename(columns=rename)

    for c in es.columns:
        if pd.api.types.is_numeric_dtype(es[c]):
            es[c] = es[c].round(2)

    return es


def _normalize_title(text: str) -> str:
    return " ".join(str(text).strip().lower().replace("ё", "е").split())


def _deduplicate_charts(charts: list[Chart]) -> list[Chart]:
    unique: list[Chart] = []
    seen: set[str] = set()

    for ch in charts:
        if not ch or not ch.title:
            continue
        key = _normalize_title(ch.title)
        if key not in seen:
            unique.append(ch)
            seen.add(key)

    return unique


def export_docx(result: AnalysisResult, charts: list[Chart], output_path: str) -> str:
    doc = Document()

    for section in doc.sections:
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    title = doc.add_heading("Отчёт по маркетинговой эффективности", level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    meta = doc.add_paragraph()
    meta.add_run(
        f"Сформирован: {datetime.now().strftime('%Y-%m-%d %H:%M')}\n"
        f"Группировка: {result.group_label}\n"
        f"Строк после очистки: {result.rows_loaded} из {result.rows_original} "
        f"(удалено {result.rows_dropped})"
    ).italic = True

    doc.add_heading("Сводка по сегментам", level=1)
    display_df = format_metrics_for_display(result)
    _add_dataframe_table(doc, display_df)

    if result.extra_summary is not None and not result.extra_summary.empty and result.extra_category_col:
        doc.add_heading(f"Доп. категория: {result.extra_category_col}", level=1)
        doc.add_paragraph(
            f"Дополнительная категория «{result.extra_category_col}» использована для "
            f"отдельной сводки по подсегментам. Ключевая метрика — "
            f"{result.metric_labels.get(result.extra_metric, result.extra_metric)}."
        )
        _add_dataframe_table(doc, _format_extra_summary(result))

    doc.add_heading("Рекомендации", level=1)
    if result.recommendations:
        for r in result.recommendations:
            doc.add_paragraph(r, style="List Bullet")
    else:
        doc.add_paragraph("Рекомендации не сформированы.")

    if result.warnings:
        doc.add_heading("Замечания и предупреждения", level=1)
        for w in result.warnings:
            doc.add_paragraph(w, style="List Bullet")

    charts = _deduplicate_charts(charts)

    if charts:
        doc.add_heading("Графики", level=1)

        for i, ch in enumerate(charts):
            if i > 0 and i % 2 == 0:
                doc.add_page_break()

            p = doc.add_paragraph()
            run = p.add_run(ch.title)
            run.bold = True

            try:
                doc.add_picture(ch.path, width=Inches(6.1))
            except Exception as e:
                doc.add_paragraph(f"(не удалось вставить график: {e})")

            if ch.description:
                p = doc.add_paragraph(ch.description)
                for run in p.runs:
                    run.italic = True
                    run.font.size = Pt(8)

    doc.save(output_path)
    return output_path